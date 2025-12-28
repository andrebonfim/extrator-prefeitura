from docx import Document
from datetime import datetime
import os

def criar_documento_teste():
  # Garante que a pasta existe
  if not os.path.exists("documentos"):
    os.makedirs("documentos")

  # Cria um novo documento
  doc = Document()

  # Define os metadados (que nosso outro script vai ler)
  prop = doc.core_properties
  prop.author = "André Luis Bonfim"
  prop.title = "Termo de Referencia - Compra de Notebooks"

  # Adiciona conteúdo ao arquivo
  doc.add_heading('Termo de Referencia n° 01/2026', 0)
  doc.add_paragraph('Este é um documento de teste criado automaticamente para o projeto de estudo.')
  doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')

  # Salva o arquivo
  nome_arquivo = "documentos/tr_notebooks_2026.docx"
  doc.save(nome_arquivo)
  print(f"Sucesso! Arquivo '{nome_arquivo}' criado com metadados.")

if __name__ == "__main__":
  criar_documento_teste()  